"""
environments/trace_env/tools/report_tool.py

Generates a DOCX report containing a table of financial transactions.
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_transactions_to_docx(transactions: list[dict], output_path: str = "financial_report.docx") -> str:
    """
    Export a list of parsed transactions into a formatted DOCX table.
    
    Args:
        transactions: List of transaction dictionaries.
        output_path: Path where the DOCX file will be saved.
        
    Returns:
        The absolute path to the generated DOCX file.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading("Financial Transactions Audit Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Total Transactions: {len(transactions)}")
    
    # Calculate total spend
    import re
    total_spend = 0.0
    for t in transactions:
        amount_str = re.sub(r'[^\d.]', '', t.get("total") or "")
        amount = float(amount_str) if amount_str else 0.0
        total_spend += amount
        
    doc.add_paragraph(f"Total Spend: ₹{total_spend:,.2f}")
    
    # Create Table
    table = doc.add_table(rows=1, cols=7) # Increased to 7 columns
    table.style = 'Table Grid'
    
    # Set headers
    hdr_cells = table.rows[0].cells
    headers = ["Date", "Vendor", "Category", "Amount", "Method", "Source", "Notes"]
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                
    # Add rows
    for t in transactions:
        row_cells = table.add_row().cells
        
        # Determine amount
        amount_str = re.sub(r'[^\d.]', '', t.get("total") or "")
        amount = float(amount_str) if amount_str else 0.0
        
        # Build notes
        details = []
        if t.get("details"):
            for k, v in t["details"].items():
                if k == "sheet_notes": continue # handle separately
                if isinstance(v, list):
                    details.append(f"{k}: {len(v)} items")
                else:
                    details.append(f"{k}: {v}")
        
        sheet_notes = t.get("notes") or t.get("details", {}).get("sheet_notes")
        if sheet_notes and str(sheet_notes) not in details:
            details.append(f"[Sheet] {sheet_notes}")

        if not details:
            subj = t.get("subject", "").strip()
            if subj:
                details.append(subj[:60])
        notes_str = " | ".join(details)
        
        source = t.get("_source", "gmail").capitalize()
        if t.get("image_analyses") or t.get("doc_analyses"):
            source += " (+AI)"

        row_cells[0].text = t.get("date", "")
        row_cells[1].text = t.get("vendor", "Unknown")
        row_cells[2].text = t.get("category", "Unknown").capitalize()
        row_cells[3].text = f"₹{amount:,.2f}" if amount > 0 else "-"
        row_cells[4].text = t.get("payment_method", "Unknown")
        row_cells[5].text = source
        row_cells[6].text = notes_str
        
    # Save the document with retry logic
    abs_path = os.path.abspath(output_path)
    try:
        doc.save(abs_path)
    except IOError:
        # Fallback if file is open
        import time
        suffix = int(time.time())
        alt_path = abs_path.replace(".docx", f"_{suffix}.docx")
        doc.save(alt_path)
        abs_path = alt_path
    
    return abs_path
