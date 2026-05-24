"""
environments/trace_env/tools/doc_tool.py

Document Extraction Pipeline for Trace — Docling-First, Local-Only.

Architecture:
─────────────────────────────────────────────────────────────────
  Stage 1 │ PARSE       │ Docling (IBM) — Universal Document Parser
          │             │ Converts PDF/DOCX/PPTX → Markdown
          │             │ Preserves layout, tables, reading order
          ▼             ▼
  Stage 2 │ IMAGES      │ Extract embedded images from document
          │             │ Forward each to image_tool (Ollama VLM)
          │             │ Moondream → qwen2-vl / llama3.2-vision
          ▼             ▼
  Stage 3 │ UNIFIED     │ Combine text + image descriptions
          │ OUTPUT      │ Standard dict → WorldModel

Fallback: If Docling is not installed, falls back to PyMuPDF (PDF only)
          or python-docx (DOCX only) for basic text extraction.

Memory strategy:
  - Docling runs locally, no GPU required
  - Peak RAM: ~2-3GB for most documents (well within 16GB)
  - Images are analysed one-at-a-time through Ollama (sequential, not parallel)

Install:
  pip install docling
  # Fallbacks (optional):
  pip install PyMuPDF python-docx
"""

from __future__ import annotations
import re
import io
import logging
import os
import tempfile
import uuid
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ── Configuration (overridden via env_config.yaml → configure()) ──────────────

_MAX_IMAGES_PER_DOC = 2          # limit VLM calls per document
_ANALYSE_EMBEDDED_IMAGES = False # DISABLED: amounts come from text layer, not logos
_MAX_TEXT_LENGTH = 50000         # truncate very long documents


def configure(config: dict):
    """
    Called once at startup with the document_extraction section of env_config.yaml.

    config = {
        "enabled": true,
        "analyse_embedded_images": true,
        "max_images_per_doc": 5,
        "max_text_length": 50000,
    }
    """
    global _MAX_IMAGES_PER_DOC, _ANALYSE_EMBEDDED_IMAGES, _MAX_TEXT_LENGTH

    _ANALYSE_EMBEDDED_IMAGES = config.get("analyse_embedded_images", _ANALYSE_EMBEDDED_IMAGES)
    _MAX_IMAGES_PER_DOC = config.get("max_images_per_doc", _MAX_IMAGES_PER_DOC)
    _MAX_TEXT_LENGTH = config.get("max_text_length", _MAX_TEXT_LENGTH)

    logger.info(
        f"[DOC_TOOL] Configured — "
        f"analyse_images={_ANALYSE_EMBEDDED_IMAGES}, "
        f"max_images={_MAX_IMAGES_PER_DOC}"
    )


# ── Public API ────────────────────────────────────────────────────────────────

def extract_document(
    file_bytes: bytes,
    filename: str,
    mime_type: str = "",
    analyse_images: bool = True,
) -> dict:
    """
    Extract text and images from a document file (PDF, DOCX, PPTX).

    Uses Docling as primary parser. Falls back to PyMuPDF (PDF) or
    python-docx (DOCX) if Docling is not installed.

    Args:
        file_bytes:     Raw document bytes (e.g. from Gmail attachment).
        filename:       Original filename (used for format detection + result ID).
        mime_type:      MIME type hint (optional, auto-detected from extension).
        analyse_images: If True, embedded images are analysed via image_tool.

    Returns:
        Unified extraction dict:
        {
            "id":               str,
            "filename":         str,
            "mime_type":        str,
            "parser_used":      "docling" | "pymupdf" | "python-docx",
            "extracted_text":   str,       # full text or Markdown
            "summary":          str,       # first ~300 chars
            "page_count":       int | None,
            "images_found":     int,
            "image_analyses":   list[dict],  # from image_tool
            "entities":         dict,        # placeholder for downstream NER
            "error":            str | None,
        }
    """
    result = _empty_result(filename, mime_type)

    try:
        ext = Path(filename).suffix.lower()

        # ── Try PyMuPDF first for PDFs (ultra-fast text extraction) ─────
        pdf_ok = False
        if ext == ".pdf":
            pdf_ok = _try_pymupdf(file_bytes, result)
            # If PyMuPDF got text, we're done — skip heavy Docling
            # If PyMuPDF returned False (empty/image-based PDF), try RapidOCR first
            if not pdf_ok:
                pdf_ok = _try_rapidocr(file_bytes, result)
            
        if not pdf_ok:
            # ── Fallback to Docling or other parsers ──────────────────
            docling_ok = _try_docling(file_bytes, filename, ext, result)

            if not docling_ok and ext in (".docx", ".doc"):
                _try_python_docx(file_bytes, result)
            elif not docling_ok and ext != ".pdf":
                result["error"] = f"Unsupported format: {ext}"
                return result

        # Truncate extremely long text
        if len(result["extracted_text"]) > _MAX_TEXT_LENGTH:
            result["extracted_text"] = result["extracted_text"][:_MAX_TEXT_LENGTH] + "\n\n[... truncated]"

        # Generate summary from extracted text
        text = result["extracted_text"]
        if text:
            result["summary"] = text[:300].strip()

        # ── Analyse embedded images via image_tool ────────────────────
        if analyse_images and _ANALYSE_EMBEDDED_IMAGES:
            _analyse_embedded_images(file_bytes, filename, ext, result)

        logger.info(
            f"[DOC_TOOL] Extraction complete for '{filename}' — "
            f"parser={result['parser_used']}, "
            f"text_len={len(result['extracted_text'])}, "
            f"images={result['images_found']}"
        )

    except Exception as e:
        err = f"{type(e).__name__}: {e}"
        result["error"] = err
        logger.error(f"[DOC_TOOL] Extraction failed for '{filename}': {err}")

    return result


# ── Docling (primary parser) ──────────────────────────────────────────────────

def _try_docling(file_bytes: bytes, filename: str, ext: str, result: dict) -> bool:
    """
    Try to parse document with Docling. Returns True if successful.
    """
    try:
        from docling.document_converter import DocumentConverter
    except ImportError:
        logger.info("[DOC_TOOL] Docling not installed, trying fallback parsers.")
        return False

    try:
        # Docling needs a file path — write to temp file
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            converter = DocumentConverter()
            doc_result = converter.convert(tmp_path)

            # Export to Markdown — preserves layout, tables, reading order
            md_content = doc_result.document.export_to_markdown()
            result["extracted_text"] = md_content
            
            # Improved regex for Indian/Global amounts
            # Matches: ₹1,500.00, Rs. 1500, Rs 1500/-, INR 1,500
            found_amounts = re.findall(r'(?:₹|Rs\.?|INR)\s*([\d,]+(?:\.\d{1,2})?)(?:\s*/-)?', md_content, re.IGNORECASE)
            if found_amounts:
                # Sort to find the highest value (likely the Total)
                sorted_amounts = sorted(found_amounts, key=lambda x: float(x.replace(',', '')), reverse=True)
                result["entities"]["amounts"] = [f"₹{a}" for a in sorted_amounts]
            
            result["parser_used"] = "docling"
            
            if result["entities"]["amounts"]:
                logger.info(f"[DOC_TOOL] Docling found {len(result['entities']['amounts'])} amounts. Top: {result['entities']['amounts'][0]}")

            # Count pages if available
            if hasattr(doc_result.document, "pages"):
                result["page_count"] = len(doc_result.document.pages)

            # Count and extract images
            if hasattr(doc_result.document, "images"):
                images = doc_result.document.images
                result["images_found"] = len(images) if images else 0

            logger.info(
                f"[DOC_TOOL] Docling parsed '{filename}' — "
                f"{len(md_content)} chars Markdown, "
                f"{result['page_count'] or '?'} pages"
            )
            return True

        finally:
            # Clean up temp file
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

    except Exception as e:
        logger.warning(f"[DOC_TOOL] Docling failed for '{filename}': {e}")
        return False


# ── Fallback: RapidOCR (image-based PDFs) ────────────────────────────────────

def _try_rapidocr(file_bytes: bytes, result: dict) -> bool:
    """
    Use RapidOCR to OCR image-based PDFs page by page.
    Much lighter than Docling's layout predictor (~200MB vs 3GB).
    Already installed: pip install rapidocr-onnxruntime
    """
    try:
        import fitz  # PyMuPDF to render pages as images
        from rapidocr import RapidOCR
    except ImportError:
        return False

    try:
        ocr = RapidOCR()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        result["page_count"] = len(doc)
        result["parser_used"] = "rapidocr"
        pages_text = []

        for page_num, page in enumerate(doc):
            # Render page to image at 150 DPI (good quality, manageable size)
            mat = fitz.Matrix(150 / 72, 150 / 72)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")

            # Run OCR
            ocr_result, _ = ocr(img_bytes)
            if ocr_result:
                page_text = "\n".join([line[1] for line in ocr_result if line and len(line) > 1])
                if page_text.strip():
                    pages_text.append(f"--- Page {page_num + 1} ---\n{page_text}")

        doc.close()
        result["extracted_text"] = "\n\n".join(pages_text)

        if not result["extracted_text"].strip():
            logger.warning("[DOC_TOOL] RapidOCR: no text found in PDF images.")
            return False

        # Pre-extract amounts
        found_amounts = re.findall(
            r'(?:[\u20b9]|Rs\.?|INR)\s*([\d,]+(?:\.\d{1,2})?)(?:\s*/-)?',
            result["extracted_text"], re.IGNORECASE
        )
        if found_amounts:
            sorted_amounts = sorted(found_amounts, key=lambda x: float(x.replace(',', '')), reverse=True)
            result["entities"]["amounts"] = [f"\u20b9{a}" for a in sorted_amounts]
            logger.info(f"[DOC_TOOL] RapidOCR found {len(sorted_amounts)} amounts. Top: \u20b9{sorted_amounts[0]}")

        logger.info(f"[DOC_TOOL] RapidOCR parsed — {len(result['extracted_text'])} chars, {result['page_count']} pages")
        return True

    except Exception as e:
        logger.warning(f"[DOC_TOOL] RapidOCR failed: {e}")
        return False


# ── Fallback: PyMuPDF (PDF only) ─────────────────────────────────────────────

def _try_pymupdf(file_bytes: bytes, result: dict) -> bool:
    """
    Fallback PDF parser using PyMuPDF (fitz). Ultra-fast text extraction.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        result["error"] = (
            "Neither Docling nor PyMuPDF installed. "
            "Install with: pip install docling  OR  pip install PyMuPDF"
        )
        return False

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        result["parser_used"] = "pymupdf"
        result["page_count"] = len(doc)

        # Extract text from all pages
        pages_text = []
        image_count = 0
        for page_num, page in enumerate(doc):
            text = page.get_text("text")
            if text.strip():
                pages_text.append(f"--- Page {page_num + 1} ---\n{text}")
            image_count += len(page.get_images(full=True))

        result["extracted_text"] = "\n\n".join(pages_text)
        result["images_found"] = image_count
        
        # If no text was extracted, PDF is likely image-based — signal failure
        # so the caller can try OCR (RapidOCR/Docling) instead.
        if not result["extracted_text"].strip():
            logger.info(f"[DOC_TOOL] PyMuPDF: no text layer found ({image_count} images). Needs OCR.")
            doc.close()
            return False
        
        # Pre-extract amounts from text to aid the verifier
        text = result["extracted_text"]
        found_amounts = re.findall(r'(?:₹|Rs\.?|INR)\s*([\d,]+(?:\.\d{1,2})?)(?:\s*/-)?', text, re.IGNORECASE)
        if found_amounts:
            sorted_amounts = sorted(found_amounts, key=lambda x: float(x.replace(',', '')), reverse=True)
            result["entities"]["amounts"] = [f"₹{a}" for a in sorted_amounts]
            logger.info(f"[DOC_TOOL] PyMuPDF found {len(sorted_amounts)} amounts. Top: ₹{sorted_amounts[0]}")
            
        doc.close()

        logger.info(
            f"[DOC_TOOL] PyMuPDF parsed — "
            f"{len(result['extracted_text'])} chars, "
            f"{result['page_count']} pages, "
            f"{image_count} images"
        )
        return True

    except Exception as e:
        result["error"] = f"PyMuPDF failed: {e}"
        logger.error(f"[DOC_TOOL] PyMuPDF error: {e}")
        return False


# ── Fallback: python-docx (DOCX only) ────────────────────────────────────────

def _try_python_docx(file_bytes: bytes, result: dict) -> bool:
    """
    Fallback DOCX parser using python-docx.
    """
    try:
        from docx import Document
    except ImportError:
        result["error"] = (
            "Neither Docling nor python-docx installed. "
            "Install with: pip install docling  OR  pip install python-docx"
        )
        return False

    try:
        doc = Document(io.BytesIO(file_bytes))
        result["parser_used"] = "python-docx"

        # Extract paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        result["extracted_text"] = "\n\n".join(paragraphs)

        # Count inline images
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
        result["images_found"] = image_count

        logger.info(
            f"[DOC_TOOL] python-docx parsed — "
            f"{len(paragraphs)} paragraphs, "
            f"{image_count} images"
        )
        return True

    except Exception as e:
        result["error"] = f"python-docx failed: {e}"
        logger.error(f"[DOC_TOOL] python-docx error: {e}")
        return False


# ── Image analysis for embedded images ────────────────────────────────────────

def _analyse_embedded_images(
    file_bytes: bytes,
    filename: str,
    ext: str,
    result: dict,
):
    """
    Extract embedded images from the document and analyse them
    with the existing image_tool (Ollama VLM pipeline).
    """
    images_to_analyse = []

    try:
        if ext == ".pdf":
            images_to_analyse = _extract_pdf_images(file_bytes)
        elif ext in (".docx", ".doc"):
            images_to_analyse = _extract_docx_images(file_bytes)
    except Exception as e:
        logger.warning(f"[DOC_TOOL] Could not extract images from '{filename}': {e}")
        return

    if not images_to_analyse:
        return

    # Limit number of images to analyse
    images_to_analyse = images_to_analyse[:_MAX_IMAGES_PER_DOC]
    logger.info(
        f"[DOC_TOOL] Analysing {len(images_to_analyse)} embedded images "
        f"from '{filename}'..."
    )

    try:
        from .image_tool import analyse_image_from_bytes
    except ImportError:
        logger.warning("[DOC_TOOL] image_tool not available for embedded image analysis.")
        return

    for i, (img_bytes, img_name) in enumerate(images_to_analyse):
        try:
            analysis = analyse_image_from_bytes(
                image_bytes=img_bytes,
                mime_type="image/png",
                filename=f"{filename}:{img_name}",
            )
            result["image_analyses"].append(analysis)
            logger.info(
                f"[DOC_TOOL] Analysed embedded image {i+1}/{len(images_to_analyse)}: "
                f"{img_name} — {len(analysis.get('extracted_text', ''))} chars"
            )
        except Exception as e:
            logger.warning(
                f"[DOC_TOOL] Image analysis failed for embedded image "
                f"'{img_name}' in '{filename}': {e}"
            )


def _extract_pdf_images(file_bytes: bytes) -> list[tuple[bytes, str]]:
    """Extract embedded images from a PDF as (bytes, name) tuples."""
    images = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page_num, page in enumerate(doc):
            for img_idx, img_info in enumerate(page.get_images(full=True)):
                xref = img_info[0]
                try:
                    base_image = doc.extract_image(xref)
                    if base_image and base_image.get("image"):
                        img_name = f"page{page_num+1}_img{img_idx+1}.{base_image.get('ext', 'png')}"
                        images.append((base_image["image"], img_name))
                except Exception:
                    continue
        doc.close()
    except ImportError:
        logger.info("[DOC_TOOL] PyMuPDF not available for PDF image extraction.")
    return images


def _extract_docx_images(file_bytes: bytes) -> list[tuple[bytes, str]]:
    """Extract embedded images from a DOCX as (bytes, name) tuples."""
    images = []
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        for rel_id, rel in doc.part.rels.items():
            if "image" in rel.reltype:
                try:
                    image_part = rel.target_part
                    img_bytes = image_part.blob
                    img_name = os.path.basename(image_part.partname)
                    images.append((img_bytes, img_name))
                except Exception:
                    continue
    except ImportError:
        logger.info("[DOC_TOOL] python-docx not available for DOCX image extraction.")
    return images


# ── Helpers ───────────────────────────────────────────────────────────────────

def _empty_result(filename: str, mime_type: str) -> dict:
    """Return a blank result skeleton."""
    return {
        "id":              f"doc_{uuid.uuid4().hex[:12]}",
        "filename":        filename,
        "mime_type":       mime_type,
        "parser_used":     "none",
        "extracted_text":  "",
        "summary":         "",
        "page_count":      None,
        "images_found":    0,
        "image_analyses":  [],
        "entities": {
            "amounts":  [],
            "dates":    [],
            "vendors":  [],
            "items":    [],
            "other":    [],
        },
        "error":           None,
    }
